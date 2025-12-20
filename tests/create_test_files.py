# -*- coding: utf-8 -*-
"""
테스트용 PDF, Word, 이미지 파일 생성
"""

from pathlib import Path

# PDF 생성
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# Word 생성
from docx import Document

# 이미지 생성
from PIL import Image, ImageDraw, ImageFont

samples_dir = Path(__file__).parent / "samples"


def create_pdf():
    """여러 페이지 PDF 생성"""
    pdf_path = samples_dir / "test_document.pdf"
    c = canvas.Canvas(str(pdf_path), pagesize=A4)
    width, height = A4

    # 페이지 1
    c.setFont("Helvetica-Bold", 24)
    c.drawString(50, height - 50, "Test Document - Page 1")
    c.setFont("Helvetica", 12)
    c.drawString(50, height - 100, "This is the first page of the test document.")
    c.drawString(50, height - 120, "Topic: Introduction to Machine Learning")
    c.drawString(50, height - 160, "Machine Learning is a subset of artificial intelligence.")
    c.drawString(50, height - 180, "It enables computers to learn from data without explicit programming.")
    c.drawString(50, height - 220, "Key concepts:")
    c.drawString(70, height - 240, "- Supervised Learning: Learning from labeled data")
    c.drawString(70, height - 260, "- Unsupervised Learning: Finding patterns in unlabeled data")
    c.drawString(70, height - 280, "- Reinforcement Learning: Learning through rewards")
    c.showPage()

    # 페이지 2
    c.setFont("Helvetica-Bold", 24)
    c.drawString(50, height - 50, "Test Document - Page 2")
    c.setFont("Helvetica", 12)
    c.drawString(50, height - 100, "Topic: Deep Learning Architecture")
    c.drawString(50, height - 140, "Neural Network Components:")
    c.drawString(70, height - 160, "1. Input Layer: Receives the raw data")
    c.drawString(70, height - 180, "2. Hidden Layers: Process and transform data")
    c.drawString(70, height - 200, "3. Output Layer: Produces the final prediction")
    c.drawString(50, height - 240, "Popular architectures include CNN, RNN, and Transformer.")
    c.drawString(50, height - 260, "The Transformer architecture is used in GPT and BERT models.")
    c.showPage()

    # 페이지 3
    c.setFont("Helvetica-Bold", 24)
    c.drawString(50, height - 50, "Test Document - Page 3")
    c.setFont("Helvetica", 12)
    c.drawString(50, height - 100, "Topic: Project Timeline")
    c.drawString(50, height - 140, "Phase 1 (Week 1-2): Data Collection")
    c.drawString(70, height - 160, "- Gather training datasets")
    c.drawString(70, height - 180, "- Clean and preprocess data")
    c.drawString(50, height - 220, "Phase 2 (Week 3-4): Model Development")
    c.drawString(70, height - 240, "- Design neural network architecture")
    c.drawString(70, height - 260, "- Implement training pipeline")
    c.drawString(50, height - 300, "Phase 3 (Week 5-6): Evaluation")
    c.drawString(70, height - 320, "- Test model performance")
    c.drawString(70, height - 340, "- Fine-tune hyperparameters")
    c.save()

    print(f"[PDF] Created: {pdf_path}")
    return pdf_path


def create_word():
    """Word 문서 생성"""
    docx_path = samples_dir / "test_document.docx"
    doc = Document()

    doc.add_heading('Cloud Computing Guide', 0)

    doc.add_heading('What is Cloud Computing?', level=1)
    doc.add_paragraph(
        'Cloud computing is the delivery of computing services over the internet. '
        'These services include servers, storage, databases, networking, software, and analytics.'
    )

    doc.add_heading('Types of Cloud Services', level=1)
    doc.add_paragraph('1. IaaS (Infrastructure as a Service)', style='List Number')
    doc.add_paragraph('   - Provides virtualized computing resources over the internet')
    doc.add_paragraph('   - Examples: AWS EC2, Google Compute Engine, Azure VMs')

    doc.add_paragraph('2. PaaS (Platform as a Service)', style='List Number')
    doc.add_paragraph('   - Provides a platform for developers to build applications')
    doc.add_paragraph('   - Examples: Heroku, Google App Engine, Azure App Service')

    doc.add_paragraph('3. SaaS (Software as a Service)', style='List Number')
    doc.add_paragraph('   - Delivers software applications over the internet')
    doc.add_paragraph('   - Examples: Gmail, Salesforce, Microsoft 365')

    doc.add_heading('Benefits of Cloud Computing', level=1)
    doc.add_paragraph('- Cost Efficiency: Pay only for what you use')
    doc.add_paragraph('- Scalability: Easily scale resources up or down')
    doc.add_paragraph('- Reliability: High availability and disaster recovery')
    doc.add_paragraph('- Security: Enterprise-grade security features')

    doc.save(str(docx_path))
    print(f"[Word] Created: {docx_path}")
    return docx_path


def create_image():
    """텍스트가 포함된 이미지 생성"""
    img_path = samples_dir / "test_image.png"

    # 800x600 흰색 배경 이미지
    img = Image.new('RGB', (800, 600), color='white')
    draw = ImageDraw.Draw(img)

    # 기본 폰트 사용 (시스템 폰트)
    try:
        font_large = ImageFont.truetype("arial.ttf", 36)
        font_medium = ImageFont.truetype("arial.ttf", 24)
        font_small = ImageFont.truetype("arial.ttf", 18)
    except:
        font_large = ImageFont.load_default()
        font_medium = font_large
        font_small = font_large

    # 제목
    draw.text((50, 30), "API Documentation", fill='black', font=font_large)

    # 내용
    draw.text((50, 100), "Endpoint: /api/v1/users", fill='navy', font=font_medium)
    draw.text((50, 140), "Method: GET", fill='black', font=font_small)
    draw.text((50, 170), "Description: Retrieves a list of all users", fill='black', font=font_small)

    draw.text((50, 220), "Endpoint: /api/v1/users/{id}", fill='navy', font=font_medium)
    draw.text((50, 260), "Method: POST", fill='black', font=font_small)
    draw.text((50, 290), "Description: Creates a new user", fill='black', font=font_small)

    draw.text((50, 340), "Authentication:", fill='navy', font=font_medium)
    draw.text((50, 380), "Bearer Token required in Authorization header", fill='black', font=font_small)

    draw.text((50, 440), "Rate Limit: 100 requests per minute", fill='darkred', font=font_medium)

    # 테두리
    draw.rectangle([10, 10, 790, 590], outline='gray', width=2)

    img.save(str(img_path))
    print(f"[Image] Created: {img_path}")
    return img_path


if __name__ == "__main__":
    print("Creating test files...")
    create_pdf()
    create_word()
    create_image()
    print("Done!")
